import docx
import re
import json

def parse_poems_docx():
    doc = docx.Document('++++Виктор  Морхат вершы (1).docx')
    
    categories = []
    current_category = None
    current_poem = []
    
    for p in doc.paragraphs:
        # Check if Heading 1
        if p.style.name.startswith('Heading 1'):
            if current_category and current_poem:
                current_category['poems'].append(current_poem)
                current_poem = []
            
            cat_title = p.text.strip()
            if cat_title:
                current_category = {'title': cat_title, 'poems': []}
                categories.append(current_category)
            continue
            
        # Check for page breaks
        has_page_break = False
        for run in p.runs:
            if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                has_page_break = True
                break
                
        if has_page_break:
            if current_category and current_poem:
                current_category['poems'].append(current_poem)
            current_poem = []
            
        # If no category yet, create a default one
        if not current_category:
            current_category = {'title': 'Лірычныя', 'poems': []}
            categories.append(current_category)
            
        text = p.text.strip()
        
        # Detect if it's italic (for epigraphs)
        is_italic = False
        if p.style.font.italic:
            is_italic = True
        else:
            # Check if majority of runs are italic
            italic_len = sum(len(r.text) for r in p.runs if r.italic)
            if italic_len > len(text) * 0.5:
                is_italic = True
                
        current_poem.append({
            'text': text,
            'is_italic': is_italic
        })
        
    if current_category and current_poem:
        current_category['poems'].append(current_poem)

    final_categories = []
    
    for cat in categories:
        parsed_poems = []
        for raw_poem in cat['poems']:
            # Filter empty lines
            lines = [l for l in raw_poem if l['text']]
            if not lines: continue
            
            title = lines[0]['text'].strip('. *-=\n\r\t')
            if title == 'Віктар Морхат.  Вершы.':
                lines = lines[1:]
                if not lines: continue
                title = lines[0]['text'].strip('. *-=\n\r\t')
                
            if not title or len(title) < 2:
                continue
                
            epigraph = []
            text_lines = []
            date = ''
            footnotes = []
            
            # Start from index 1 (after title)
            idx = 1
            
            # Detect epigraph (italic lines immediately after title)
            while idx < len(lines) and lines[idx]['is_italic'] and not re.search(r'^\* \* \*', lines[idx]['text']):
                epigraph.append(lines[idx]['text'])
                idx += 1
                
            # Process remaining lines
            remaining_lines = lines[idx:]
            
            for line_obj in remaining_lines:
                txt = line_obj['text']
                
                # Remove asterisks
                if re.match(r'^[\*\s\-\=]{3,}$', txt):
                    continue
                    
                # Detect date
                if re.match(r'^(19|20)\d{2}\s*г\.?$', txt) or re.match(r'^\d{4}$', txt) or ' г.' in txt[-5:]:
                    date = txt
                    continue
                    
                # Detect footnote (often long lines containing dashes or parenthesis, or appearing after the date)
                # We can also check if the line is very long (not a typical poem stanza line) and has a dash
                if len(txt) > 60 and (' – ' in txt or ' - ' in txt or '(' in txt):
                    footnotes.append(txt)
                    continue
                    
                if txt != title and txt != title + '.':
                    text_lines.append(txt)
                    
            # Reconstruct stanzas based on empty lines in original logic (Wait, docx loses empty paragraphs if we stripped them)
            # Let's reconstruct using simple joining, but docx paragraphs preserve natural flow.
            # We didn't keep empty paragraphs in `lines = [l for l in raw_poem if l['text']]`. 
            # Let's fix this: we SHOULD keep empty paragraphs to reconstruct stanzas!
            pass # We will fix this by redefining the loop below

# Redoing with stanza preservation
def parse_poems_docx2():
    doc = docx.Document('++++Виктор  Морхат вершы (1).docx')
    
    categories = []
    current_category = None
    current_poem = []
    
    for p in doc.paragraphs:
        if p.style.name.startswith('Heading 1'):
            if current_category and any(l['text'].strip() for l in current_poem):
                current_category['poems'].append(current_poem)
            current_poem = []
            cat_title = p.text.strip()
            if cat_title:
                current_category = {'title': cat_title, 'poems': []}
                categories.append(current_category)
            continue
            
        has_page_break = False
        for run in p.runs:
            if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                has_page_break = True
                break
                
        if has_page_break:
            if current_category and any(l['text'].strip() for l in current_poem):
                current_category['poems'].append(current_poem)
            current_poem = []
            
        if not current_category:
            current_category = {'title': 'Лірычныя', 'poems': []}
            categories.append(current_category)
            
        text = p.text.strip()
        
        is_italic = False
        if p.style.font.italic:
            is_italic = True
        else:
            italic_len = sum(len(r.text) for r in p.runs if r.italic)
            if len(text) > 0 and italic_len > len(text) * 0.5:
                is_italic = True
                
        # Check for indent
        indent = p.paragraph_format.left_indent
        is_indented = indent is not None and 0 < indent < 1000000
        
        # Preserve empty lines for stanzas
        for line_text in text.split('\n'):
            line_text = line_text.replace('\t', '    ').strip()
            # If replacing tab leaves multiple commas or floating commas, fix it
            line_text = re.sub(r',\s*,', ',', line_text)
            current_poem.append({
                'text': line_text,
                'is_italic': is_italic,
                'is_empty': len(line_text) == 0,
                'is_indented': is_indented
            })
        
    if current_category and any(l['text'].strip() for l in current_poem):
        current_category['poems'].append(current_poem)

    final_categories = []
    
    seen_titles = {}
    
    for cat in categories:
        parsed_poems = []
        for raw_poem in cat['poems']:
            # Strip leading empty lines
            while raw_poem and raw_poem[0]['is_empty']:
                raw_poem.pop(0)
            if not raw_poem: continue
            
            title = raw_poem[0]['text'].strip('. *-=\n\r\t')
            if 'Віктар Морхат' in title and 'Вершы' in title:
                raw_poem.pop(0)
                while raw_poem and raw_poem[0]['is_empty']:
                    raw_poem.pop(0)
                if not raw_poem: continue
                title = raw_poem[0]['text'].strip('. *-=\n\r\t')
                
            if not title or len(title) < 2:
                continue
                
            epigraph = []
            text_lines = []
            date = ''
            footnotes = []
            
            idx = 1
            # Skip empty lines after title
            while idx < len(raw_poem) and raw_poem[idx]['is_empty']:
                idx += 1
                
            # Extract epigraph
            while idx < len(raw_poem) and raw_poem[idx]['is_italic'] and not raw_poem[idx]['is_empty']:
                if raw_poem[idx]['text'].strip() in ['1', '2', '3']:
                    break
                epigraph.append(raw_poem[idx]['text'])
                idx += 1
                
            # Extract date (look at the end)
            end_idx = len(raw_poem) - 1
            while end_idx > idx and (raw_poem[end_idx]['is_empty'] or re.match(r'^[\*\s\-\=]{3,}$', raw_poem[end_idx]['text'])):
                end_idx -= 1
                
            # Extract footnotes from the end
            while end_idx > idx:
                txt = raw_poem[end_idx]['text']
                if len(txt) > 40 and (' – ' in txt or ' - ' in txt or '(лац' in txt.lower()):
                    footnotes.insert(0, txt)
                    end_idx -= 1
                elif raw_poem[end_idx]['is_empty'] or re.match(r'^[\*\s\-\=]{3,}$', txt):
                    end_idx -= 1
                else:
                    break
                    
            # Skip empty lines/asterisks before date
            while end_idx > idx and (raw_poem[end_idx]['is_empty'] or re.match(r'^[\*\s\-\=]{3,}$', raw_poem[end_idx]['text'])):
                end_idx -= 1
                
            # Extract date from the end
            if end_idx > idx:
                txt = raw_poem[end_idx]['text']
                if re.search(r'\b(19|20)\d{2}\b', txt) or ' г.' in txt:
                    txt_clean = re.sub(r'[\*\s]+$', '', txt)
                    if '    ' in txt_clean:
                        parts = re.split(r'\s{4,}', txt_clean)
                        date = parts[-1]
                        raw_poem[end_idx]['text'] = txt[:txt.rfind(date)].strip()
                    else:
                        date = txt_clean
                        end_idx -= 1
                    
            # Extract main text
            last_was_indented = None
            for i in range(idx, end_idx + 1):
                p_data = raw_poem[i]
                txt = p_data['text']
                
                if re.match(r'^[\*\s\-\=]{3,}$', txt):
                    continue # skip asterisks
                    
                if p_data['is_empty']:
                    text_lines.append('')
                    last_was_indented = None
                else:
                    if txt != title and txt != title + '.':
                        text_lines.append(txt)
                        
                    last_was_indented = p_data['is_indented']
                        
            # Isolate part numbers
            iso_lines = []
            for txt in text_lines:
                if txt.strip() in ['1', '2', '3']:
                    if iso_lines and iso_lines[-1] != '':
                        iso_lines.append('')
                    iso_lines.append(txt.strip())
                    iso_lines.append('')
                else:
                    iso_lines.append(txt)
            text_lines = iso_lines
                        
            # Clean up trailing empty lines
            while text_lines and text_lines[-1] == '':
                text_lines.pop()
                
            # Remove purely dotted lines
            text_lines = [l for l in text_lines if not re.match(r'^[\.\s]+$', l)]
            
            # Find footnotes
            footnote_start = -1
            for i in range(len(text_lines)):
                line = text_lines[i]
                if line.startswith('*') and len(line.replace('*', '').strip()) > 2:
                    # check if this is near the end
                    if i > len(text_lines) - 20:
                        footnote_start = i
                        break
                        
            if footnote_start != -1:
                f_lines = text_lines[footnote_start:]
                combined = []
                for line in f_lines:
                    if line == '':
                        continue
                    if line.startswith('*'):
                        combined.append(line)
                    else:
                        if combined:
                            combined[-1] += ' ' + line
                        else:
                            combined.append(line)
                            
                footnote_text = '\n'.join(combined)
                footnotes.append(footnote_text)
                
                text_lines = text_lines[:footnote_start]
                while text_lines and text_lines[-1] == '':
                    text_lines.pop()
                        
            # Auto-chunking for missing empty lines
            chunked_lines = []
            current_block = []
            
            for line in text_lines:
                if line == '':
                    if len(current_block) >= 8:
                        for i in range(0, len(current_block), 4):
                            chunked_lines.extend(current_block[i:i+4])
                            if i + 4 < len(current_block):
                                chunked_lines.append('')
                    else:
                        chunked_lines.extend(current_block)
                    chunked_lines.append('')
                    current_block = []
                else:
                    current_block.append(line)
            
            if current_block:
                if len(current_block) >= 8:
                    for i in range(0, len(current_block), 4):
                        chunked_lines.extend(current_block[i:i+4])
                        if i + 4 < len(current_block):
                            chunked_lines.append('')
                else:
                    chunked_lines.extend(current_block)
                    
            text_lines = chunked_lines
            
            # Force alternating indentation for poem 16
            if 'Роздум перад сустрэчай' in title:
                formatted_lines = []
                stanza_idx = 0
                for line in text_lines:
                    if line == '':
                        formatted_lines.append('')
                        stanza_idx += 1
                    else:
                        clean_line = line.replace('\t', '')
                        if stanza_idx % 2 == 1:
                            formatted_lines.append('\t' + clean_line)
                        else:
                            formatted_lines.append(clean_line)
                text_lines = formatted_lines
                        
            # Clean up text_lines
            while text_lines and text_lines[-1] == '':
                text_lines.pop()
                
            # Fix typo in Poem 2 (author's style: single line)
            text_lines = [l.replace('халады…Над', 'халады… Над') for l in text_lines]
                
            text_clean = '\n'.join(text_lines)
            
            original_title = title
            if original_title in seen_titles:
                seen_titles[original_title] += 1
                title = f"{original_title} (Варыянт {seen_titles[original_title]})"
            else:
                seen_titles[original_title] = 1
                
            subtitle_val = ''
            
            # Check if subtitle got sucked into epigraph (if it was italicized)
            if epigraph:
                first_epi_line = epigraph[0].strip('. ').lower()
                if first_epi_line.startswith('санет') or first_epi_line.startswith('ода') or first_epi_line.startswith('трыпціх') or first_epi_line.startswith('элегія') or first_epi_line.startswith('альбо ') or first_epi_line.startswith('або '):
                    subtitle_val = epigraph.pop(0)
                    
            epigraph_val = '\n'.join(epigraph)
            
            # Extract common subtitle formats from main text (if not italicized)
            first_line_end = text_clean.find('\n')
            first_line = text_clean[:first_line_end].strip() if first_line_end != -1 else text_clean.strip()
            fl_lower = first_line.lower().strip('. ')
            
            if fl_lower.startswith('санет') or fl_lower.startswith('элегія') or fl_lower.startswith('альбо ') or fl_lower.startswith('або '):
                if not subtitle_val:
                    subtitle_val = first_line
                if first_line_end != -1:
                    text_clean = text_clean[first_line_end:].strip()
                else:
                    text_clean = ''
                    
            # Custom fix for "Творчасць" epigraph SECOND
            if original_title == 'Творчасць':
                if 'Бродскага.' in text_clean:
                    parts = text_clean.split('Бродскага.', 1)
                    epigraph_val = (parts[0] + 'Бродскага.').strip()
                    text_clean = parts[1].strip()
                    
            # Custom fix for "Поль Верлен" epigraph
            if 'Поль Верлен' in text_clean:
                parts = text_clean.split('Поль Верлен', 1)
                # Only if it's at the beginning of the text
                if len(parts[0]) < 100:
                    new_epigraph = (parts[0] + 'Поль Верлен').strip()
                    if epigraph_val:
                        epigraph_val = epigraph_val + '\n\n' + new_epigraph
                    else:
                        epigraph_val = new_epigraph
                    text_clean = parts[1].strip()
            
            # Clean up trailing spaces/commas in title
            title = title.rstrip(' ,')
            
            # Custom rename for Poem 5
            if title == 'Гэй, хто з намі ?!':
                title = 'Гэй, хто з намі ?! (Філасофскі варыянт)'
            
            parsed_poems.append({
                'title': title,
                'subtitle': subtitle_val,
                'epigraph': epigraph_val,
                'text': text_clean,
                'date': date,
                'footnotes': '\n'.join(footnotes)
            })
            
        if parsed_poems:
            final_categories.append({
                'title': cat['title'],
                'poems': parsed_poems
            })
            
    # Save
    js_content = "const categories = " + json.dumps(final_categories, ensure_ascii=False, indent=4) + ";\n"
    
    with open('v4_final_concept/data.js', 'w', encoding='utf-8') as f:
        f.write(js_content)

if __name__ == '__main__':
    parse_poems_docx2()
